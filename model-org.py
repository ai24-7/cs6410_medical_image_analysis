from docx import Document # type: ignore

# Function to create a Word document with the model summary
def create_word_summary(summary_text):
    # Create a new Word document
    doc = Document()
    
    # Add a title
    doc.add_heading('Model Summary', level=1)
    
    # Add the model summary text
    doc.add_paragraph(summary_text)
    
    # Save the document
    doc.save('model_summary.docx')
    print("Model summary saved as model_summary.docx")

# Model summary text
model_summary_text = """
Layer (type)               Output Shape         Param #
===========================================================
            Conv2d-1         [-1, 64, 112, 112]           9,408
       BatchNorm2d-2         [-1, 64, 112, 112]             128
              ReLU-3         [-1, 64, 112, 112]               0
         MaxPool2d-4           [-1, 64, 56, 56]               0
            Conv2d-5           [-1, 64, 56, 56]          36,864
       BatchNorm2d-6           [-1, 64, 56, 56]             128
              ReLU-7           [-1, 64, 56, 56]               0
            Conv2d-8           [-1, 64, 56, 56]          36,864
       BatchNorm2d-9           [-1, 64, 56, 56]             128
             ReLU-10           [-1, 64, 56, 56]               0
       BasicBlock-11           [-1, 64, 56, 56]               0
           Conv2d-12           [-1, 64, 56, 56]          36,864
      BatchNorm2d-13           [-1, 64, 56, 56]             128
             ReLU-14           [-1, 64, 56, 56]               0
           Conv2d-15           [-1, 64, 56, 56]          36,864
      BatchNorm2d-16           [-1, 64, 56, 56]             128
             ReLU-17           [-1, 64, 56, 56]               0
       BasicBlock-18           [-1, 64, 56, 56]               0
           Conv2d-19           [-1, 64, 56, 56]          36,864
      BatchNorm2d-20           [-1, 64, 56, 56]             128
             ReLU-21           [-1, 64, 56, 56]               0
           Conv2d-22           [-1, 64, 56, 56]          36,864
      BatchNorm2d-23           [-1, 64, 56, 56]             128
             ReLU-24           [-1, 64, 56, 56]               0
       BasicBlock-25           [-1, 64, 56, 56]               0
           Conv2d-26          [-1, 128, 28, 28]          73,728
      BatchNorm2d-27          [-1, 128, 28, 28]             256
             ReLU-28          [-1, 128, 28, 28]               0
           Conv2d-29          [-1, 128, 28, 28]         147,456
      BatchNorm2d-30          [-1, 128, 28, 28]             256
           Conv2d-31          [-1, 128, 28, 28]           8,192
      BatchNorm2d-32          [-1, 128, 28, 28]             256
"""

# Create Word document with the model summary
create_word_summary(model_summary_text)
